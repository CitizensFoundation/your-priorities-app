from docx.api import Document
import json
import re
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from docx.document import Document as doctwo

def iter_block_items(parent):
    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

#document = Document("/home/robert/Documents/WorldBank/Surveys/refinalthingstofixfor1_roundofsurveys/Comm Rep_Final_ENG_clear_June 22 2020.docx")
#document = Document("/home/robert/Documents/WorldBank/Surveys/refinalthingstofixfor1_roundofsurveys/Comm Rep_Final_RU_clear_June 22 2020.docx")
#document = Document("/home/robert/Documents/WorldBank/Surveys/refinalthingstofixfor1_roundofsurveys/Comm Rep_Final_KGZ_clear_June 22 2020.docx")

#document = Document("/home/robert/Documents/WorldBank/Surveys/refinalthingstofixfor1_roundofsurveys/Govt Officials_Final_ENG_clear_June 22 2020.docx")
#document = Document("/home/robert/Documents/WorldBank/Surveys/refinalthingstofixfor1_roundofsurveys/Govt Officials_Final_RU_clear_June 22 2020.docx")
document = Document("/home/robert/Documents/WorldBank/Surveys/refinalthingstofixfor1_roundofsurveys/Govt Officials_Final_KGZ_clear_June 22 2020.docx")

survey_items = []

def hasNumbers(inputString):
  return any(char.isdigit() for char in inputString)

def hasBrackets(inputString):
  return any(char=="[" for char in inputString)

def getStringToBracket(inputString):
  return inputString
  #return inputString.split("[")[0].strip()

def reallyGetStringToBracket(inputString):
  return inputString.split("[")[0].strip()

def appendSurveyItem(item):
  print(item)
  survey_items.append(item)

def clean(text):
  return text.strip()

def has_specify(text):
  if text.find("___")>-1 or text.find("(specify)")>-1 or text.find('─┴─┴─')>-1 or text.find("(указать)")>-1 or text.find("(белгилоо)")>-1:
    return True
  else:
    return False

def appendRatio(uniqueId, optionsText, questionText, subType=None):
  print("RADIOS")
  radios = []
  uniqueId = uniqueId.replace("]","").replace("]","")
  if len(optionsText)<2:
    print(optionsText)
    optionsText = splitByDigit("".join(optionsText)).split("\n")
    print(optionsText)

  for option in optionsText:
    number = option[:2].strip().replace(".","")
    #print("Number: "+number)
    text = option[2:].strip()
    #print("Text: "+text)
    isSpecify = False
    skipTo = None
    if has_specify(text):
      isSpecify = True
      text = text.replace("_","")
      if len(text)==0:
        text = "____"
    text, skipTo = processSkipTo(text)
    sub_type = get_sub_type(text, text, text)
    text = text.replace(",","")
    text = text.replace(":",";")
    text = text.replace("─┴","")
    text = text.replace("─┘", "")
    text = text.strip()
    text = text.replace("└", " ")
    if len(text)>0 and len(number)>0:
      radios.append({'skipTo': skipTo, 'text': text, 'subCode': number, 'subType': sub_type, 'isSpecify': isSpecify})
  questionText, max_length = get_max_length(questionText)
  print("MAX:"+str(max_length))
  appendSurveyItem({'type':'radios', 'subType': subType, 'uniqueId': uniqueId, 'radioButtons': radios, 'text': getStringToBracket(questionText), 'maxLength': max_length})

def processSkipTo(text):
  skipTo = None
  if text.find("skip to")>-1:
      skipTo = text.split("skip to")[1].strip()
      text = reallyGetStringToBracket(text.split("skip to")[0]).strip()
  if text.find("перейти к")>-1:
      skipTo = text.split("перейти к")[1].strip()
      text = reallyGetStringToBracket(text.split("перейти к")[0]).strip()
  if text.find("өтүү")>-1:
      skipTo = text.split("өтүү")[1].strip()
      text = reallyGetStringToBracket(text.split("өтүү")[0]).strip()

  return text, skipTo

def get_max_length(text):
  max_length = 100
  if text.find("<")>-1 and text.find(">")>-1:
    max_length = text[text.find("<")+1:text.rfind(">")]
    text = text.replace("<"+max_length+">","")

  return text,max_length

def appendCheckbox(uniqueId, optionsText, questionText):
  print("CHECKBOXES")
  checkboxes = []
  uniqueId = uniqueId.replace("]","").replace("]","")
  for option in optionsText:
    number = option[:2].strip().replace(".","")
    text = option[2:].strip()
    isSpecify = False
    skipTo = None
    if has_specify(text):
      isSpecify = True
      text = text.replace("_","")
      if len(text)==0:
        text = "____"
    text, skipTo = processSkipTo(text)
    sub_type = get_sub_type(text, text, text)
    text = text.replace(",","")
    text = text.replace(":",";")
    text = text.strip()
    if len(text)>0 and len(number)>0:
      checkboxes.append({'skipTo': skipTo, 'text': text, 'subCode': number, 'isSpecify': isSpecify,'subType': sub_type})

  questionText, max_length = get_max_length(questionText)
  appendSurveyItem({'type':'checkboxes', 'uniqueId': uniqueId, 'checkboxes': checkboxes, 'text': getStringToBracket(questionText), 'maxLength': max_length})

def splitByDigit(text):
  text = text.strip()
  part = ""
  for char in text:
    if char.isdigit():
      part+="\n"
    part += char
  part+="\n"
  print("PART: "+part)
  return part.strip()

def splitByColumns(columns):
  print(columns)
  i = 0
  options = []
  for column in columns:
    if i>1:
      options.append(column.text)
    i += 1
  print(options)
  return options

firstCell = True

def get_sub_type(text, descriptionOne, descriptionTwo):
  sub_type = "text"

  match = re.search(r'.{2}/.{2}/.{4}', text)

  if match:
    sub_type = "date"
  elif descriptionOne.find('─┴─┴─')>-1 or descriptionTwo.find('─┴─┴─')>-1:
    sub_type = "number"

  return sub_type

for block in iter_block_items(document):
    if isinstance(block, Paragraph):
      if len(block.text.strip())>0:
        appendSurveyItem({'type':'textHeader','text': block.text.strip()})
    if isinstance(block, Table):
      table = block
      denseUniqueId = None
      for ri, row in enumerate(table.rows):
        if len(table.rows)>ri+1:
          print("HEH: "+table.rows[ri+1].cells[0].text.strip())
          print("TRUE"+str(table.rows[ri+1].cells[0].text.strip().startswith("а")))
        if firstCell==True:
            firstCell=False
            print("Description")
            appendSurveyItem({'type':'textDescription','text': getStringToBracket(row.cells[0].text.strip())})
        else: #if row.cells[2] and len(row.cells[2].text)>1
          print("ROWS"+str(len(table.rows))+" ri: "+str(ri))
          if len(table.rows)>ri+1 and (table.rows[ri+1].cells[0].text.strip().startswith("a") or table.rows[ri+1].cells[0].text.strip().startswith("а")):
            denseUniqueId = row.cells[0].text.strip().replace("[","").replace("]","")
            print("DENSE RATIOS: "+denseUniqueId)
            appendSurveyItem({'type':'textDescription','text': getStringToBracket(row.cells[1].text.strip())})
            #if len(row.cells)>2 and row.cells[2].text.strip().startswith("1"):
              #appendSurveyItem({'type':'textDescription','text': getStringToBracket(row.cells[2].text.strip())})

          elif len(row.cells)>3 and (row.cells[3].text.strip().startswith("A ") or row.cells[3].text.strip().startswith("A_") or row.cells[3].text.strip().startswith("А ")):
            appendCheckbox(row.cells[0].text.strip(), row.cells[3].text.split("\n"), row.cells[1].text.strip())
          elif row.cells[2].text.strip().startswith("A") or row.cells[2].text.strip().startswith("A_") or row.cells[2].text.strip().startswith("А "):
            appendCheckbox(row.cells[0].text.strip(), row.cells[2].text.split("\n"), row.cells[1].text.strip())

          elif len(row.cells[0].text.strip())==1 and len(row.cells)>4 and row.cells[2].text.strip().startswith("1"):
            appendRatio(denseUniqueId+row.cells[0].text.strip(), splitByColumns(row.cells), row.cells[1].text.strip(),'rating')

          elif len(row.cells[0].text.strip())==1 and len(row.cells)>3 and len(row.cells[3].text)>5:
            appendRatio(denseUniqueId+row.cells[0].text.strip(), splitByDigit(row.cells[3].text).split("\n"), row.cells[1].text.strip(), 'rating')
          elif len(row.cells[0].text.strip())==1 and len(row.cells)>2 and len(row.cells[2].text)>5:
            appendRatio(denseUniqueId+row.cells[0].text.strip(), splitByDigit(row.cells[2].text).split("\n"), row.cells[1].text.strip(), 'rating')

          elif len(row.cells)>3 and (row.cells[3].text.strip().startswith("1 ") or row.cells[3].text.strip().startswith("1. ")):
            appendRatio(row.cells[0].text.strip(), row.cells[3].text.split("\n"), row.cells[2].text.strip())
          elif row.cells[2].text.strip().startswith("1") or row.cells[2].text.strip().startswith("1."):
            appendRatio(row.cells[0].text.strip(), row.cells[2].text.split("\n"), row.cells[1].text.strip())

          elif row.cells[0] and len(row.cells[0].text)>3 and len(row.cells[0].text)<6:
            text = getStringToBracket(row.cells[1].text.strip())
            max_length = 100
            text, max_length = get_max_length(text)

            sub_type = get_sub_type(text, row.cells[1].text.strip(), row.cells[2].text.strip())

            if len(getStringToBracket(row.cells[1].text.strip()))<24:
              appendSurveyItem({'type':'textFieldLong', 'subType': sub_type, 'uniqueId': row.cells[0].text.strip(), 'maxLength': max_length, 'text': text })
            else:
              appendSurveyItem({'type':'textFieldLong', 'subType': sub_type, 'uniqueId': row.cells[0].text.strip(), 'maxLength': max_length, 'text': text })

        #else:
         # print("NOTHING")

print(json.dumps(survey_items))
